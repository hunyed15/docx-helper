# -*- coding: utf-8 -*-
"""docx-helper 回归测试套件。

运行：
    python -m unittest discover -s tests -v
或：
    python tests/test_format.py -v
"""
import os
import re
import sys
import tempfile
import unittest
import zipfile

# 允许从技能根目录 import scripts/format.py
SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(SKILL_DIR, 'scripts'))

import format as fmt  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt  # noqa: E402


def make_docx(paragraphs, with_numpr=False, with_image=False, dest=None):
    """构造测试 docx。paragraphs: 文本列表。返回文件路径。"""
    if dest is None:
        fd, dest = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
    doc = Document()
    for text in paragraphs:
        p = doc.add_paragraph(text)
        if with_numpr:
            # 模拟自动编号：numId=1（中文序列）
            ppr = p._p.get_or_add_pPr()
            numpr = ppr.makeelement(qn('w:numPr'), {})
            ilvl = ppr.makeelement(qn('w:ilvl'), {qn('w:val'): '0'})
            numid = ppr.makeelement(qn('w:numId'), {qn('w:val'): '1'})
            numpr.append(ilvl)
            numpr.append(numid)
            ppr.append(numpr)
    if with_image:
        # 模拟浮动图片 anchor
        p = doc.add_paragraph()
        p._p.append(_make_floating_anchor())
    doc.save(dest)
    return dest


def _make_floating_anchor():
    """构造一个浮动图片 anchor 元素（behindDoc=1 + wrapThrough）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsmap, nsdecls
    anchor = OxmlElement('wp:anchor')
    anchor.set('behindDoc', '1')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    # wrapThrough
    wrap = OxmlElement('wp:wrapThrough')
    anchor.append(wrap)
    return anchor


class TestResetFormat(unittest.TestCase):
    def test_reset_keeps_text(self):
        src = make_docx(['第一章 总则', '第一条 本规范适用于公共数据平台。', '二〇二六年八月'])
        dst = src + '.reset.docx'
        try:
            fmt.reset_format(src, dst)
            doc = Document(dst)
            texts = [p.text for p in doc.paragraphs]
            self.assertEqual(texts[0], '第一章 总则')
            self.assertEqual(texts[1], '第一条 本规范适用于公共数据平台。')
            self.assertEqual(texts[2], '二〇二六年八月')
        finally:
            for f in (src, dst):
                if os.path.exists(f):
                    os.remove(f)


class TestNumberBaking(unittest.TestCase):
    def test_format_num_chinese(self):
        """中文序列格式应正确转换（之前 japaneseCounting/chineseCounting 缺失的回归）。"""
        self.assertEqual(fmt._format_num(1, 'chineseCounting'), '一')
        self.assertEqual(fmt._format_num(2, 'chineseCounting'), '二')
        self.assertEqual(fmt._format_num(3, 'japaneseCounting'), '三')
        self.assertEqual(fmt._format_num(10, 'japaneseCounting'), '十')
        self.assertEqual(fmt._format_num(11, 'japaneseCounting'), '十一')

    def test_format_num_decimal(self):
        self.assertEqual(fmt._format_num(5, 'decimal'), '5')
        self.assertEqual(fmt._format_num(1, 'decimal'), '1')

    def test_reset_keeps_numbered_text(self):
        """reset 应保留文字内容（含已烘焙编号文本）。"""
        src = make_docx(['一、总则', '（一）编制目的', '1.适用范围'])
        dst = src + '.reset.docx'
        try:
            fmt.reset_format(src, dst)
            doc = Document(dst)
            texts = [p.text for p in doc.paragraphs]
            self.assertEqual(texts[0], '一、总则')
            self.assertEqual(texts[1], '（一）编制目的')
            self.assertEqual(texts[2], '1.适用范围')
        finally:
            for f in (src, dst):
                if os.path.exists(f):
                    os.remove(f)


class TestFormatters(unittest.TestCase):
    def test_signature_format(self):
        doc = Document()
        p = doc.add_paragraph('巡检日期：    年  月  日')
        fmt.format_signature(p)
        self.assertEqual(p.alignment, 0)  # LEFT
        self.assertEqual(p.paragraph_format.first_line_indent, Pt(0))
        self.assertIsNotNone(p.paragraph_format.left_indent)

    def test_caption_format(self):
        doc = Document()
        p = doc.add_paragraph('图 2-1 应急事件处置组织架构图')
        fmt.format_caption(p)
        self.assertEqual(p.alignment, 1)  # CENTER
        if p.runs:
            rpr = p.runs[0]._element.rPr
            ea = rpr.rFonts.get(qn('w:eastAsia')) if rpr is not None and rpr.rFonts is not None else None
            self.assertIn('黑体', ea or '')

    def test_title_format(self):
        doc = Document()
        p = doc.add_paragraph('宿迁市公共数据平台数据开放管理规范')
        fmt.format_title(p)
        self.assertEqual(p.alignment, 1)  # CENTER


class TestCheckContent(unittest.TestCase):
    def test_typo_detection(self):
        src = make_docx(['第一条 日常日常巡检', '（2）防范未然', '（4）一但发生故障'])
        try:
            report = fmt.check_content(src)
            found = [x['found'] for x in report['typos']]
            self.assertIn('日常日常', found)
            self.assertIn('防范未然', found)
            self.assertIn('一但', found)
        finally:
            if os.path.exists(src):
                os.remove(src)

    def test_numbering_jump_detection(self):
        src = make_docx(['（1）第一项', '（2）第二项', '（4）第四项（跳号）'])
        try:
            report = fmt.check_content(src)
            issues = [x['issue'] for x in report['numbering_issues']]
            self.assertTrue(any('跳号' in i for i in issues))
        finally:
            if os.path.exists(src):
                os.remove(src)

    def test_user_typo_rules(self):
        """用户自定义错别字库应合并生效。"""
        cfg = {'typo_rules': {'布署': '部署'}}
        fmt._load_user_typo_rules(cfg)
        src = make_docx(['平台已完成布署'])
        try:
            report = fmt.check_content(src)
            found = [x['found'] for x in report['typos']]
            self.assertIn('布署', found)
        finally:
            if os.path.exists(src):
                os.remove(src)
            fmt._load_user_typo_rules(fmt.DEFAULT_CONFIG)  # 还原


class TestFixImages(unittest.TestCase):
    def test_floating_to_inline(self):
        src = make_docx(['第一章 总则', '正文内容'], with_image=True)
        try:
            doc = Document(src)
            n = fmt._fix_floating_images(doc)
            self.assertEqual(n, 1)
            doc.save(src)  # 保存回文件再验证
            with zipfile.ZipFile(src) as z:
                xml = z.read('word/document.xml').decode('utf-8')
                self.assertEqual(len(re.findall(r'<wp:anchor', xml)), 0)
                self.assertGreaterEqual(len(re.findall(r'<wp:inline', xml)), 1)
        finally:
            import gc
            gc.collect()
            for _ in range(3):
                try:
                    os.remove(src)
                    break
                except PermissionError:
                    import time
                    time.sleep(0.3)


class TestValidateInput(unittest.TestCase):
    def test_nonexistent_file(self):
        ok, err = fmt.validate_input('/nonexistent/path.docx')
        self.assertFalse(ok)
        self.assertIn('找不到', err)

    def test_doc_extension(self):
        fd, path = tempfile.mkstemp(suffix='.doc')
        os.close(fd)
        try:
            ok, err = fmt.validate_input(path)
            self.assertFalse(ok)
            self.assertIn('仅支持 .docx', err)
        finally:
            os.remove(path)

    def test_valid_docx(self):
        src = make_docx(['正文'])
        try:
            ok, err = fmt.validate_input(src)
            self.assertTrue(ok)
        finally:
            if os.path.exists(src):
                os.remove(src)

    def test_corrupt_file(self):
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.write(fd, b'not a zip file at all')
        os.close(fd)
        try:
            ok, err = fmt.validate_input(path)
            self.assertFalse(ok)
        finally:
            os.remove(path)


class TestTemplate(unittest.TestCase):
    def test_new_document_from_template(self):
        fd, out = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        os.remove(out)
        try:
            result = fmt.new_document_from_template(out, title='测试规范')
            self.assertTrue(os.path.exists(result))
            doc = Document(result)
            self.assertGreater(len(doc.paragraphs), 0)
        finally:
            if os.path.exists(out):
                os.remove(out)


class TestApplyFormat(unittest.TestCase):
    """apply_format 端到端测试：验证排版后页面/字体/页码正确。"""

    def test_apply_produces_valid_docx(self):
        """apply 应输出可打开的 docx，段落不丢失。"""
        src = make_docx(['第一章 总则', '1.1 编制目的', '正文内容测试'])
        dst = src + '.v1.docx'
        try:
            result = fmt.apply_format(src, dst)
            self.assertTrue(os.path.exists(result))
            doc = Document(result)
            self.assertEqual(len(doc.paragraphs), 3)
            # 页面边距应已设置
            sec = doc.sections[0]
            self.assertAlmostEqual(sec.top_margin.cm, 3.7, places=1)
        finally:
            for f in (src, dst):
                if os.path.exists(f):
                    os.remove(f)

    def test_apply_with_structure(self):
        """带 structure 映射排版后字体应正确。"""
        src = make_docx(['总则', '编制目的', '正文'])
        dst = src + '.v1.docx'
        struct = {"paragraphs": {"0": "title", "1": "section", "2": "body"}, "cover": False, "title_index": 0}
        struct_path = src + '.json'
        import json
        with open(struct_path, 'w', encoding='utf-8') as f:
            json.dump(struct, f)
        try:
            result = fmt.apply_format(src, dst, structure=struct)
            doc = Document(result)
            # 段0 应居中（title）
            self.assertEqual(doc.paragraphs[0].alignment, 1)  # CENTER
        finally:
            for f in (src, dst, struct_path):
                if os.path.exists(f):
                    os.remove(f)


class TestFixImagesInTable(unittest.TestCase):
    """表格内浮动图片修复测试。"""

    def test_fix_images_in_table_cell(self):
        """表格单元格内的浮动图片也应被转为 inline。"""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        doc = Document()
        # 添加表格，在单元格里放浮动图片
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        anchor = OxmlElement('wp:anchor')
        anchor.set('behindDoc', '1')
        anchor.set('simplePos', '0')
        p._p.append(anchor)
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        doc.save(path)
        try:
            n = fmt._fix_floating_images(Document(path))
            self.assertEqual(n, 1, "表格内浮动图片应被修复")
        finally:
            os.remove(path)


class TestCheckContentNumberingEdgeCases(unittest.TestCase):
    """编号检测边界用例。"""

    def test_numbering_dup_detection(self):
        """同级编号重复应检出。"""
        src = make_docx(['（1）第一项', '（1）重复项'])
        try:
            report = fmt.check_content(src)
            issues = [x['issue'] for x in report['numbering_issues']]
            self.assertTrue(any('重复' in i for i in issues))
        finally:
            if os.path.exists(src):
                os.remove(src)


if __name__ == '__main__':
    unittest.main(verbosity=2)
