#!/usr/bin/env python3
"""
文档排版助手 (docx-helper) — 对已有 .docx 文档套用规范排版。

推荐工作流（大模型驱动）:
  0. 重置模式: python format.py --reset <input.docx>
     → 彻底清空所有手动格式（字体/字号/颜色/加粗/缩进/自动编号），
        输出 {原名}+reset.docx，得到一个"中性"文档，便于下一步判结构。

  1. 大模型判结构: 阅读 reset 文档全文，决定每个段落的类型
        （title/h1/h2/h3/h4/bullet/body），写入 structure.json:
        {"paragraphs": {"1": "h2", "3": "h2", "5": "h4"}, "cover": false, "title_index": null}

  2. 应用模式: python format.py --apply <reset.docx> --structure structure.json
     → 按结构映射套用排版，输出版本化文件: {原名}+docx-helper+v{N}.docx

  3. 直接模式（兼容，无 structure 时走启发式）: python format.py <input.docx>
     → 等同于 --apply（自动分配版本号）

版本化命名:
  - 第一次运行: 报告+docx-helper+v1.docx
  - 第二次运行: 报告+docx-helper+v2.docx
  - 每次自动递增版本号，永不覆盖原文件或旧版本

排版规则（字体/字号/对齐）:
  1. 页面: A4, 上 3.7 / 下 3.5 / 左 2.7 / 右 2.7 cm
  2. title     (大标题):    2号 方正小标宋_GBK, 居中
  3. chapter   (章标题):    3号 方正黑体_GBK, 左对齐
  4. section   (节标题):    3号 方正楷体_GBK, 左对齐
  5. subsection(子节标题):  3号 方正仿宋_GBK, 左对齐
  6. item      (条目编号):  3号 方正仿宋_GBK, 左对齐, 无缩进
  7. bullet    (项目符号):  3号 方正仿宋_GBK, 左对齐
  8. body      (正文):      3号 方正仿宋_GBK, 首行缩进 2 字, 行距 28.9 磅
  9. 页码: 4号 Times New Roman, "— N —" 格式, 居中

注意：标题层级判定应按语义深度（relative depth），而非按编号字符正则匹配。
例：1.编制背景 和 2.1 盘点目标 同为「章的直接子级」→ 均判 section（楷体）；4.1.1 按分层迁移 是「h2 的子级」→ 判 subsection（仿宋）。
旧名 h1/h2/h3/h4 仍可接收（format.py 内部自动别名映射）。

字体依赖: 需安装方正小标宋/黑体/楷体/仿宋_GBK, 否则 Word 会提示字体缺失。
"""

import glob, itertools, json, os, re, sys
from collections import defaultdict, Counter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ══════════════════════════════════════════════════════════════════════════════
# 排版参数（可通过 .docx-helper.json 配置覆盖）
# ══════════════════════════════════════════════════════════════════════════════

BLACK = RGBColor(0, 0, 0)

# ── 运行时变量（由 _apply_config 设置，不要直接硬编码） ──
FONT_TITLE = '方正小标宋_GBK'
FONT_CHAPTER = '方正黑体_GBK'
FONT_SECTION = '方正楷体_GBK'
FONT_SUBSECTION = '方正仿宋_GBK'
FONT_ITEM = '方正仿宋_GBK'
FONT_BODY = '方正仿宋_GBK'
FONT_EN = 'Times New Roman'
# 向下兼容别名
FONT_H1 = FONT_CHAPTER; FONT_H2 = FONT_SECTION; FONT_H3 = FONT_SUBSECTION; FONT_H4 = FONT_ITEM

SIZE_TITLE = Pt(22)
SIZE_CHAPTER = Pt(16)
SIZE_SECTION = Pt(16)
SIZE_SUBSECTION = Pt(16)
SIZE_ITEM = Pt(16)
SIZE_BODY = Pt(16)
SIZE_PAGE = Pt(14)
SIZE_2 = SIZE_TITLE; SIZE_3 = SIZE_BODY; SIZE_4 = SIZE_PAGE

MARGINS = {}
FOOTER_DIST = Cm(2.54)
LINE_SPACING = Pt(28.9)
BODY_INDENT_CHARS = 2
PAGE_NUMBER_FORMAT = '\u2014 N \u2014'

# ── 默认配置（GB/T 9704-2012 规范） ──
DEFAULT_CONFIG = {
    "page": {"margins": {"top": "3.7cm", "bottom": "3.5cm", "left": "2.7cm", "right": "2.7cm"}, "footer_distance": "2.54cm"},
    "fonts": {"title": "方正小标宋_GBK", "chapter": "方正黑体_GBK", "section": "方正楷体_GBK", "subsection": "方正仿宋_GBK", "item": "方正仿宋_GBK", "body": "方正仿宋_GBK", "en": "Times New Roman"},
    "sizes": {"title": 22, "chapter": 16, "section": 16, "subsection": 16, "item": 16, "body": 16, "page_number": 14},
    "spacing": {"line_height": 28.9, "body_indent_chars": 2},
    "page_number_format": "\u2014 N \u2014",
}

def _parse_cm(val):
    if isinstance(val, (int, float)): return Cm(float(val))
    s = str(val).strip()
    if s.endswith('mm'): return Cm(float(s.replace('mm','')) / 10)
    return Cm(float(s.replace('cm','')))

def _deep_merge(base, override):
    for k, v in override.items():
        if k in base and isinstance(base[k], dict) and isinstance(v, dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v

def _apply_config(cfg):
    global MARGINS, FOOTER_DIST, LINE_SPACING, BODY_INDENT_CHARS, PAGE_NUMBER_FORMAT
    global FONT_TITLE, FONT_CHAPTER, FONT_SECTION, FONT_SUBSECTION, FONT_ITEM, FONT_BODY, FONT_EN
    global FONT_H1, FONT_H2, FONT_H3, FONT_H4
    global SIZE_TITLE, SIZE_CHAPTER, SIZE_SECTION, SIZE_SUBSECTION, SIZE_ITEM, SIZE_BODY, SIZE_PAGE
    global SIZE_2, SIZE_3, SIZE_4
    p = cfg.get("page", {}); m = p.get("margins", {})
    MARGINS = dict(top=_parse_cm(m.get("top","3.7cm")), bottom=_parse_cm(m.get("bottom","3.5cm")),
                   left=_parse_cm(m.get("left","2.7cm")), right=_parse_cm(m.get("right","2.7cm")))
    FOOTER_DIST = _parse_cm(p.get("footer_distance", "2.54cm"))
    f = cfg.get("fonts", {})
    FONT_TITLE = f.get("title", "方正小标宋_GBK")
    FONT_CHAPTER = f.get("chapter", "方正黑体_GBK")
    FONT_SECTION = f.get("section", "方正楷体_GBK")
    FONT_SUBSECTION = f.get("subsection", "方正仿宋_GBK")
    FONT_ITEM = f.get("item", "方正仿宋_GBK")
    FONT_BODY = f.get("body", "方正仿宋_GBK")
    FONT_EN = f.get("en", "Times New Roman")
    FONT_H1 = FONT_CHAPTER; FONT_H2 = FONT_SECTION; FONT_H3 = FONT_SUBSECTION; FONT_H4 = FONT_ITEM
    s = cfg.get("sizes", {})
    SIZE_TITLE = Pt(s.get("title", 22)); SIZE_CHAPTER = Pt(s.get("chapter", 16))
    SIZE_SECTION = Pt(s.get("section", 16)); SIZE_SUBSECTION = Pt(s.get("subsection", 16))
    SIZE_ITEM = Pt(s.get("item", 16)); SIZE_BODY = Pt(s.get("body", 16))
    SIZE_PAGE = Pt(s.get("page_number", 14))
    SIZE_2 = SIZE_TITLE; SIZE_3 = SIZE_BODY; SIZE_4 = SIZE_PAGE
    sp = cfg.get("spacing", {})
    LINE_SPACING = Pt(sp.get("line_height", 28.9))
    BODY_INDENT_CHARS = sp.get("body_indent_chars", 2)
    PAGE_NUMBER_FORMAT = cfg.get("page_number_format", "\u2014 N \u2014")
    _rebuild_type_meta()

def load_config(path=None):
    cfg = json.loads(json.dumps(DEFAULT_CONFIG))
    if path and os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as fh:
            _deep_merge(cfg, json.load(fh))
    _apply_config(cfg)
    return cfg

TYPE_META = {}

def _rebuild_type_meta():
    TYPE_META.clear()
    TYPE_META.update({
        'title': ('大标题', FONT_TITLE, SIZE_TITLE, 'center'),
        'chapter': ('章标题', FONT_CHAPTER, SIZE_CHAPTER, 'left'),
        'section': ('节标题', FONT_SECTION, SIZE_SECTION, 'left'),
        'subsection': ('子节标题', FONT_SUBSECTION, SIZE_SUBSECTION, 'left'),
        'item': ('条目编号', FONT_ITEM, SIZE_ITEM, 'left'),
        'bullet': ('项目符号', FONT_BODY, SIZE_BODY, 'left'),
        'body': ('正文', FONT_BODY, SIZE_BODY, 'left'),
        'h1': ('一级标题', FONT_CHAPTER, SIZE_CHAPTER, 'left'),
        'h2': ('二级标题', FONT_SECTION, SIZE_SECTION, 'left'),
        'h3': ('三级标题', FONT_SUBSECTION, SIZE_SUBSECTION, 'left'),
        'h4': ('四级标题', FONT_ITEM, SIZE_ITEM, 'left'),
    })

# ══════════════════════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════════════════════

_apply_config(DEFAULT_CONFIG)  # 模块导入时初始化所有排版常量

def _set_font(run, cn_font, en_font, size):
    """设置 run 的字体、字号，确保不加粗、颜色为黑色"""
    run.font.size = size
    run.bold = False
    run.font.color.rgb = BLACK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
        rFonts.set(qn(f'w:{attr}'), cn_font if attr == 'eastAsia' else en_font)


def _set_spacing(para, space_before=0, space_after=0):
    pf = para.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = LINE_SPACING


def _add_page_numbers(section):
    """在默认页脚添加页码（居中，格式由 PAGE_NUMBER_FORMAT 配置）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = PAGE_NUMBER_FORMAT
    parts = fmt.split('N', 1)
    if len(parts) == 2:
        if parts[0]:
            _add_page_run(p, parts[0])
        _add_page_field(p)
        if parts[1]:
            _add_page_run(p, parts[1])
    else:
        _add_page_run(p, '— ')
        _add_page_field(p)
        _add_page_run(p, ' —')


def _add_page_run(para, text):
    run = para.add_run(text)
    run.font.size = SIZE_PAGE
    run.bold = False
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_EN}" w:ascii="{FONT_EN}" w:hAnsi="{FONT_EN}"/>')
    rPr.insert(0, rFonts)


def _add_page_field(para):
    run = para.add_run()
    run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = para.add_run()
    r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r3 = para.add_run()
    r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


# ══════════════════════════════════════════════════════════════════════════════
# 版本号管理
# ══════════════════════════════════════════════════════════════════════════════

def _next_version(input_path):
    base = os.path.splitext(os.path.basename(input_path))[0]
    m = re.match(r'^(.+?)\+docx-helper\+v(\d+)$', base)
    if m:
        base = m.group(1)
    dirname = os.path.dirname(input_path) or '.'
    pattern = os.path.join(dirname, f'{base}+docx-helper+v*.docx')
    existing = glob.glob(pattern)
    max_v = 0
    for f in existing:
        m2 = re.search(r'\+docx-helper\+v(\d+)\.docx$', f)
        if m2:
            max_v = max(max_v, int(m2.group(1)))
    return base, max_v + 1, os.path.join(dirname, f'{base}+docx-helper+v{max_v + 1}.docx')


# ══════════════════════════════════════════════════════════════════════════════
# 多层级编号识别
# ══════════════════════════════════════════════════════════════════════════════

# 编号匹配模式 (按优先级从高到低, 避免短模式吃掉长模式)
_NUMBERING_PATTERNS = [
    # 多级数字编号: 4.1.1 / 4.1 / 1.2.3 等
    (re.compile(r'^(\d+(?:\.\d+)+)\s'),            'h3_multi'),
    # 中文数字一级: 一、二、三、...
    (re.compile(r'^[一二三四五六七八九十]+[、]'),       'h1'),
    # 中文数字二级: （一）（二）（三）...
    (re.compile(r'^（[一二三四五六七八九十]+）'),        'h2'),
    # 阿拉伯数字三级: 1. 2. 3. ...
    (re.compile(r'^(\d+)[\.、]\s'),                 'h3'),
    # 阿拉伯数字三级变体: 1）2）3）...
    (re.compile(r'^(\d+)）'),                       'h3'),
    # 括号四级: (1) (2) (3) ...
    (re.compile(r'^\((\d+)\)'),                     'h4'),
    # 圈号四级: ①②③④⑤⑥⑦⑧⑨⑩
    (re.compile(r'^[①-⑩]'),                         'h4'),
    # 小写字母: a. b. c. ... 或 a）b）...
    (re.compile(r'^([a-z])[\.）]\s'),               'h5'),
    # 大写字母: A. B. C. ...
    (re.compile(r'^([A-Z])[\.）]\s'),               'h5'),
    # 罗马数字: I. II. III. ...
    (re.compile(r'^[IVX]+[\.、]\s'),                'h3'),
    # 项目符号: ● ◆ ★ ■ ▶ ▪ ○ □ △
    (re.compile(r'^[●◆★■▶▪○□△▷◇♦☞✓✔]'),           'bullet'),
    # 第X章/第X节
    (re.compile(r'^第[一二三四五六七八九十\d]+[章节]'),   'chapter'),
    # 附件
    (re.compile(r'^附件[一二三四五六七八九十\d]*[:：]?$'), 'h1'),
]


def _detect_numbering(text):
    """识别段落开头的编号类型, 返回 (type, match_text, rest_text) 或 None"""
    for pattern, ptype in _NUMBERING_PATTERNS:
        m = pattern.match(text)
        if m:
            match_text = m.group()
            rest = text[m.end():].strip()
            return ptype, match_text, rest
    return None


def _classify(para):
    """返回 (type, confidence, numbering_info)"""
    text = para.text.strip()
    if not text:
        return 'body', 'high', None

    style = (para.style.name if para.style else '')

    # Word 样式信任
    if 'Heading 4' in style or '标题 4' in style:
        return 'h4', 'high', None
    if 'Heading 3' in style or '标题 3' in style:
        return 'h3', 'high', None
    if 'Heading 2' in style or '标题 2' in style:
        return 'h2', 'high', None
    if 'Heading 1' in style or '标题 1' in style:
        return 'h1', 'high', None
    if 'Title' in style or '标题' in style:
        return 'title', 'high', None

    # 编号识别
    result = _detect_numbering(text)
    if result:
        ptype, match_text, rest = result
        if ptype == 'chapter':
            return 'h1', 'high', result
        if ptype == 'h5':
            return 'h4', 'high', result
        if ptype == 'h3_multi':
            return 'h3', 'high', result
        return ptype, 'high', result

    return 'body', 'high', None


def _find_title_idx(paragraphs):
    """找到文档中第一个应该作为大标题的段落索引"""
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ''
        if 'Title' in style or '标题' in style:
            return i
        if i <= 4 and len(text) <= 80:
            return i
        break
    return -1


# 编号检测返回的类型中，h3_multi / chapter / h5 需要映射到现有 TYPE_META
# 同时提供新旧别名：chapter=h1, section=h2, subsection=h3, item=h4
_TYPE_ALIAS = {
    'h3_multi': 'h3',
    'chapter': 'h1',
    'h5': 'h4',
    # 新名 → 旧名（统一收敛到 formatter）
    'chapter':    'h1',
    'section':    'h2',
    'subsection': 'h3',
    'item':       'h4',
}

def _resolve_type(ptype):
    return _TYPE_ALIAS.get(ptype, ptype)


def _strip_numpr(para):
    """移除段落的自动编号（w:numPr），使编号数字由正文文本承载，避免数字与正文字体不一致。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)


def _clear_run_formatting(run):
    """清空 run 的全部手动格式（字体/字号/颜色/加粗/斜体/下划线/字距等），仅保留文字。"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        for tag in ('w:rFonts', 'w:b', 'w:i', 'w:color', 'w:sz', 'w:u',
                    'w:highlight', 'w:spacing', 'w:w', 'w:shd', 'w:vertAlign',
                    'w:rStyle', 'w:strike', 'w:emboss', 'w:imprint', 'w:outline',
                    'w:dstrike', 'w:effect'):
            for child in rPr.findall(qn(tag)):
                rPr.remove(child)
    run.font.size = None
    run.bold = None
    run.italic = None
    run.underline = None
    if run.font.color is not None:
        run.font.color.rgb = None


# ══════════════════════════════════════════════════════════════════════════════
# 自动编号 → 文字：reset 时把 Word 自动编号(numPr)烘焙成正文文本，
# 否则删除 numPr 会把自动生成的序号（1. / 一、 / (1)）一起删掉，层级就丢了。
# ══════════════════════════════════════════════════════════════════════════════

_CN_NUM = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
_CIRCLED = {1: '①', 2: '②', 3: '③', 4: '④', 5: '⑤', 6: '⑥', 7: '⑦', 8: '⑧',
            9: '⑨', 10: '⑩', 11: '⑪', 12: '⑫', 13: '⑬', 14: '⑭', 15: '⑮',
            16: '⑯', 17: '⑰', 18: '⑱', 19: '⑲', 20: '⑳'}

def _to_chinese(n):
    if n <= 10:
        return _CN_NUM[n]
    if n < 20:
        return '十' + (_CN_NUM[n - 10] if n > 10 else '')
    if n < 100:
        t, o = divmod(n, 10)
        return _CN_NUM[t] + '十' + (_CN_NUM[o] if o else '')
    # 100–999
    if n < 1000:
        h, r = divmod(n, 100)
        result = _CN_NUM[h] + '百'
        if r == 0:
            return result
        if r < 10:
            return result + '零' + _CN_NUM[r]
        return result + _to_chinese(r)
    return str(n)

def _to_roman(n, upper=True):
    vals = [(1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'), (100, 'C'),
            (90, 'XC'), (50, 'L'), (40, 'XL'), (10, 'X'), (9, 'IX'),
            (5, 'V'), (4, 'IV'), (1, 'I')]
    result = ''
    for v, r in vals:
        while n >= v:
            result += r
            n -= v
    return result if upper else result.lower()

def _to_letters(n, upper=True):
    result = ''
    while n > 0:
        n -= 1
        result = chr(ord('A' if upper else 'a') + n % 26) + result
        n //= 26
    return result

def _format_num(n, fmt):
    if fmt in ('chineseCount', 'chineseNumber', 'ideographTraditional', 'ideographZodiac', 'chineseLegalSimplified'):
        return _to_chinese(n)
    if fmt in ('decimalEnclosedCircle', 'decimalEnclosedCircleChinese'):
        return _CIRCLED.get(n, f'({n})')
    if fmt == 'decimalFullWidth':
        return ''.join(chr(ord('０') + int(c)) for c in str(n))
    if fmt == 'decimalHalfWidth':
        return ''.join(chr(ord(' ') + int(c)) for c in str(n))
    if fmt == 'upperRoman':
        return _to_roman(n, upper=True)
    if fmt == 'lowerRoman':
        return _to_roman(n, upper=False)
    if fmt == 'upperLetter':
        return _to_letters(n, upper=True)
    if fmt == 'lowerLetter':
        return _to_letters(n, upper=False)
    return str(n)  # decimal 及未知格式回退为阿拉伯数字

def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _build_numbering_maps(doc):
    """返回 (numId→abstractNumId, abstractNumId→{ilvl:信息})"""
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return {}, {}
    if numbering_part is None:
        return {}, {}
    root = numbering_part._element
    num_to_abs = {}
    lvl_info = {}
    for num in root.findall(qn('w:num')):
        nid = num.get(qn('w:numId'))
        absn = num.find(qn('w:abstractNumId'))
        if absn is not None:
            num_to_abs[nid] = absn.get(qn('w:val'))
    for absn in root.findall(qn('w:abstractNum')):
        aid = absn.get(qn('w:abstractNumId'))
        lvl_info[aid] = {}
        for lvl in absn.findall(qn('w:lvl')):
            ilvl = int(lvl.get(qn('w:ilvl')))
            fmt = lvl.find(qn('w:numFmt'))
            txt = lvl.find(qn('w:lvlText'))
            start = lvl.find(qn('w:start'))
            lvl_info[aid][ilvl] = {
                'fmt': fmt.get(qn('w:val')) if fmt is not None else 'decimal',
                'text': txt.get(qn('w:val')) if txt is not None else '%1.',
                'start': int(start.get(qn('w:val'))) if start is not None else 1,
            }
    return num_to_abs, lvl_info

def _compute_list_texts(doc):
    """返回 {段落 lxml 元素: 计算出的编号文本}"""
    num_to_abs, lvl_info = _build_numbering_maps(doc)
    if not num_to_abs:
        return {}
    result = {}
    counters = {}      # (numId, ilvl) -> 当前序号
    seq_decimal = 0    # 顶层十进制 "1." 标题的连续序号
    prev_meta = None   # (numId, ilvl, fmt) 上一个编号段落
    parent_chain = []  # 父级链条，用于 %1.%2 嵌套编号
    for p in _iter_all_paragraphs(doc):
        numPr = p._element.find(qn('w:pPr'))
        np = numPr.find(qn('w:numPr')) if numPr is not None else None
        if np is None:
            # 非编号段落不重置计数器（正文夹在列表项中间时，编号应延续）
            continue
        numId = np.find(qn('w:numId'))
        ilvl = np.find(qn('w:ilvl'))
        if numId is None:
            prev_meta = None
            continue
        nid = numId.get(qn('w:val'))
        il = int(ilvl.get(qn('w:val')) if ilvl is not None else 0)
        absid = num_to_abs.get(nid)
        info = lvl_info.get(absid, {}).get(il) if absid is not None else None
        if info is None:
            prev_meta = (nid, il, '?')
            continue
        fmt = info['fmt']
        lvltext = info['text']
        # 顶层十进制 "1." / "1、" 标题：跨独立列表顺次重排（修复"每段都显示 1."的坏源）
        is_top_decimal = (il == 0 and fmt == 'decimal' and re.match(r'^%1[.、]?$', lvltext))
        if is_top_decimal:
            cont = (prev_meta is not None and prev_meta[1] == 0 and prev_meta[2] == 'decimal'
                    and _numid_gt(nid, prev_meta[0]))
            seq_decimal = (seq_decimal + 1) if cont else info['start']
            txt = re.sub(r'%1', _format_num(seq_decimal, fmt), lvltext)
        else:
            # 其余（嵌套、括号、(1)、项目符号等）用各自 (numId, ilvl) 计数器
            if prev_meta is None or prev_meta[0] != nid:
                for k in list(counters):
                    if k[0] == nid:
                        del counters[k]
                counters[(nid, il)] = info['start']
            elif il == prev_meta[1]:
                counters[(nid, il)] = counters.get((nid, il), info['start'] - 1) + 1
            else:
                counters[(nid, il)] = counters.get((nid, il), info['start'] - 1) + 1
            val = counters[(nid, il)]
            # 维护父级链条以支持 %1.%2 嵌套编号
            parent_chain = [pc for pc in parent_chain if pc[1] <= il]
            parent_chain.append((nid, il, val))
            chain_sorted = sorted(parent_chain, key=lambda x: x[1])
            def _rep(m, _fmt=fmt):
                k = int(m.group(1)); idx = k - 1
                if idx < len(chain_sorted):
                    return _format_num(chain_sorted[idx][2], _fmt)
                return ''
            txt = re.sub(r'%(\d)', _rep, lvltext)
        if re.search(r'[.、）\)]$', txt):
            txt += ' '
        result[p._element] = txt
        prev_meta = (nid, il, fmt)
    return result


def _numid_gt(a, b):
    """numId 通常为创建顺序的整数，后建的列表 id 更大"""
    try:
        return int(a) > int(b)
    except ValueError:
        return a > b

def _convert_list_numbers_to_text(doc):
    """把自动编号烘焙为正文文本（在段首插入编号 run），并移除 numPr。"""
    texts = _compute_list_texts(doc)
    if not texts:
        return
    for p in _iter_all_paragraphs(doc):
        el = p._element
        if el not in texts:
            continue
        num_text = texts[el]
        cur = p.text.strip()
        if cur and cur.startswith(num_text.strip()):
            _strip_numpr(p)
            continue
        run = p.add_run(num_text)
        pPr = el.find(qn('w:pPr'))
        if pPr is not None:
            pPr.addnext(run._r)
        else:
            el.insert(0, run._r)
        _strip_numpr(p)

def reset_format(input_path, output_path):
    """彻底重置文档的所有手动格式，输出了一个中性文档。

    - 段落样式置为 Normal，清除对齐/间距/行距/缩进/分页符
    - 移除自动编号 w:numPr（让编号数字回到正文文本里）
    - 清空每个 run 的字体/字号/颜色/加粗等手动格式
    - 表格内文字同样处理
    不改变文字内容，不改变段落数量与顺序（便于按索引套用结构映射）。
    """
    doc = Document(input_path)
    normal_style = None
    try:
        normal_style = doc.styles['Normal']
    except KeyError:
        normal_style = None

    # 先把自动编号转成文字（避免删除 numPr 时把序号一起删掉）
    _convert_list_numbers_to_text(doc)

    def _reset_para(para):
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = None
        pf.first_line_indent = Pt(0)
        pf.page_break_before = False
        _strip_numpr(para)
        if normal_style is not None:
            try:
                para.style = normal_style
            except Exception:
                pass
        for r in para.runs:
            _clear_run_formatting(r)

    for para in doc.paragraphs:
        _reset_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _reset_para(para)

    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# 封面检测与分析（材料/方案/报告）
# ══════════════════════════════════════════════════════════════════════════════

# 封面排版惯例参数
COVER_SPACE_ABOVE_TITLE = 7       # 标题上方空行数（约页面 1/3 处）
COVER_SPACE_BELOW_TITLE = 5       # 标题区与署名区之间空行数
COVER_SPACE_BETWEEN_ELEMENTS = 0  # 同区内元素间不空行（同行距）

_DATE_PAT = re.compile(r'^(20\d{2})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日$|^(20\d{2})[-./年](\d{1,2})[-./月](\d{1,2})')


def _detect_cover(paragraphs):
    """检测文档首页是否为材料/方案/报告的封面页。

    返回 None 或 cover_info dict:
      {
        'has_cover': True,
        'title_indices': [0],         # 大标题段落索引
        'subtitle_indices': [1],      # 副标题段落索引（可能为空列表）
        'attribution_start': 5,       # 署名区起始段落索引
        'unit_indices': [5],          # 单位名称段落索引
        'date_indices': [6],          # 日期段落索引
        'end_index': 6,               # 封面结束段落索引（之后加 page break）
        'confidence': 'high',         # 检测置信度
        'suggestions': [...],         # 格式建议
      }
    """
    paras = [(i, p) for i, p in enumerate(paragraphs) if p.text.strip()]
    if len(paras) < 3:
        return None

    # 封面特征：前 N 个段落短、居中、无编号
    COVER_WINDOW = 10  # 在前 10 个有效段落中检测

    candidates = []
    for i, para in paras[:COVER_WINDOW]:
        text = para.text.strip()
        align = str(para.alignment) if para.alignment else 'NONE'
        has_numbering = _detect_numbering(text) is not None
        candidates.append({
            'index': i,
            'text': text,
            'len': len(text),
            'align': align,
            'has_numbering': has_numbering,
            'is_date': bool(_DATE_PAT.match(text)),
        })

    # 如果前几个段落有公文编号（一、），不是封面
    early_numbered = sum(1 for c in candidates[:3] if c['has_numbering'])
    if early_numbered >= 2:
        return None

    # 封面特征评分
    score = 0
    short_count = sum(1 for c in candidates if c['len'] <= 60)
    centered_count = sum(1 for c in candidates if 'CENTER' in c['align'])
    date_count = sum(1 for c in candidates if c['is_date'])

    score += min(short_count, 5) * 1
    score += centered_count * 2
    score += date_count * 3
    score -= early_numbered * 3

    if score < 4:
        return None

    # ── 识别封面各元素 ──
    title_indices = []
    subtitle_indices = []
    unit_indices = []
    date_indices = []

    non_empty = [(i, paragraphs[i].text.strip(), paragraphs[i])
                 for i, p in enumerate(paragraphs) if p.text.strip()]

    # 找标题（封面第一个短文本，不含编号和日期）
    for idx, text, para in non_empty:
        if idx > 10:
            break
        if _detect_numbering(text):
            break  # 遇到编号段落，封面结束
        if _DATE_PAT.match(text):
            break
        if len(text) <= 80:
            if not title_indices:
                title_indices.append(idx)
            elif len(title_indices) == 1 and len(text) <= 60:
                # 第二个短文本，且紧跟在标题后
                if idx - title_indices[-1] <= 2:
                    subtitle_indices.append(idx)
                else:
                    break
            else:
                break

    if not title_indices:
        return None

    # 在封面区域内找单位和日期
    cover_start = title_indices[0]
    cover_window_end = min(cover_start + 15, len(paragraphs))

    for idx in range(cover_start + 1, cover_window_end):
        para = paragraphs[idx]
        text = para.text.strip()
        if not text:
            continue
        # 跳过已识别的标题和副标题
        if idx in title_indices or idx in subtitle_indices:
            continue
        if _detect_numbering(text):
            break  # 正文开始
        if _DATE_PAT.match(text):
            date_indices.append(idx)
            break  # 日期通常是封面最后一个元素
        if 4 <= len(text) <= 80:
            if not unit_indices:
                unit_indices.append(idx)

    # 封面结束位置
    end_index = date_indices[-1] if date_indices else (unit_indices[-1] if unit_indices else (subtitle_indices[-1] if subtitle_indices else title_indices[-1]))

    confidence = 'high' if date_count > 0 else 'medium'

    # 生成建议
    suggestions = []
    if subtitle_indices:
        suggestions.append({
            'para_index': subtitle_indices[0],
            'element': '副标题',
            'suggestion': '副标题三号方正楷体_GBK，居中，主标题正下方同行距，可前加破折号"——"',
            'severity': 'info',
        })
    if unit_indices:
        suggestions.append({
            'para_index': unit_indices[0],
            'element': '单位名称',
            'suggestion': f'单位名称三号方正仿宋_GBK，居中，与标题区间空{COVER_SPACE_BELOW_TITLE}行',
            'severity': 'info',
        })
    if date_indices:
        suggestions.append({
            'para_index': date_indices[0],
            'element': '日期',
            'suggestion': '日期三号方正仿宋_GBK，居中，单位名称正下方同行距',
            'severity': 'info',
        })
    suggestions.append({
        'para_index': title_indices[0],
        'element': '大标题',
        'suggestion': f'大标题二号方正小标宋_GBK，居中，页面上方空{COVER_SPACE_ABOVE_TITLE}行（约页面1/3处）',
        'severity': 'info',
    })

    return {
        'has_cover': True,
        'title_indices': title_indices,
        'subtitle_indices': subtitle_indices,
        'unit_indices': unit_indices,
        'date_indices': date_indices,
        'end_index': end_index,
        'confidence': confidence,
        'score': score,
        'suggestions': suggestions,
    }

# GB/T 9704-2012 §7.3.3 规定的标准四级编号体系
GB_STANDARD = {
    'name': 'GB/T 9704-2012',
    'reference': '第 7.3.3 条',
    'principle': '结构层次序数依次用"一、"→"（一）"→"1."→"（1）"',
    'chain': ['h1', 'h2', 'h3', 'h4'],  # 标准编号顺序
    'fonts': {'h1': '黑体', 'h2': '楷体', 'h3': '仿宋', 'h4': '仿宋'},
    'patterns': {
        'h1': '一、二、三、...',
        'h2': '（一）（二）（三）...',
        'h3': '1. 2. 3. ...',
        'h4': '（1）（2）（3）...',
    },
    'description': (
        '第一层用黑体字、第二层用楷体字、第三层和第四层用仿宋体字。'
        '超出四层时，国标未作规定，可参照惯例使用①、a. 等。'
    ),
}


def _analyze_numbering_hierarchy(paragraphs):
    """扫描全文，推断编号层级结构，与 GB/T 9704-2012 对比。

    返回:
      {
        'levels': [...],
        'hierarchy_chain': '一级 → 二级 → ...',
        'gb_standard': { ... },           # 国标规范
        'gb_comparison': { ... },         # 与国标的差异
        'suggestions': [                  # 逐段建议
          {'para_index': 15, 'numbering': '●', 'issue': '项目符号不在国标内',
           'suggestion': '建议改为"（N）"格式的第四级编号，如无特殊需要可删除符号直接作为正文'},
        ],
        'issues': [...],
      }
    """
    level_usage = defaultdict(list)
    seen_order = []
    all_items = []

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        result = _detect_numbering(text)
        if result:
            ptype, match_text, rest = result
            if ptype in ('h1', 'h2', 'h3', 'h4', 'h5', 'h3_multi', 'bullet', 'chapter'):
                normalized = 'h1' if ptype == 'chapter' else ('h4' if ptype == 'h5' else ptype)
                if normalized not in seen_order:
                    seen_order.append(normalized)
                level_usage[normalized].append({
                    'index': i,
                    'numbering': match_text,
                    'preview': text[:50],
                })
            all_items.append((i, ptype, match_text, text[:50]))

    # 构建层级描述
    levels = []
    for t in seen_order:
        items = level_usage[t]
        labels = {
            'h1': '中文数字+顿号',
            'h2': '括号中文数字',
            'h3': '阿拉伯数字+点',
            'h3_multi': '多级数字编号',
            'h4': '括号数字/圈号/字母',
            'bullet': '项目符号',
        }
        levels.append({
            'type': t,
            'label': labels.get(t, t),
            'count': len(items),
            'examples': [it['preview'] for it in items[:3]],
        })

    # ── 与 GB/T 9704-2012 对比 ────────────────────────────────────────────
    gb_chain = GB_STANDARD['chain']
    doc_norm = [t for t in seen_order if t in ('h1', 'h2', 'h3', 'h4', 'h3_multi', 'bullet')]
    # 将 h3_multi 映射为 h3 同层
    doc_semantic = []
    for t in doc_norm:
        doc_semantic.append('h3' if t == 'h3_multi' else t)

    gb_comparison = {
        'standard_chain': ['一、', '（一）', '1.', '（1）'],
        'document_chain': [_type_label(t) for t in doc_norm],
        'extra_types': [],      # 国标外的编号类型
        'missing_types': [],    # 文档缺失的国标层级
        'order_consistent': True,
        'notes': [],
    }

    # 检测文档中不在国标规范内的编号
    NON_STANDARD_TYPES = {'bullet', 'h3_multi', 'h5'}
    for t in doc_norm:
        if t in NON_STANDARD_TYPES:
            gb_comparison['extra_types'].append(t)

    # 检测文档中缺失的国标层级
    for gb_t in gb_chain:
        if gb_t not in doc_norm and gb_t not in doc_semantic:
            gb_comparison['missing_types'].append(gb_t)

    # 生成对比说明
    if gb_comparison['extra_types'] or gb_comparison['missing_types']:
        gb_comparison['order_consistent'] = False
        parts = []
        if gb_comparison['extra_types']:
            for t in gb_comparison['extra_types']:
                label = _type_label(t)
                parts.append(f'使用了国标未规定的编号格式"{label}"')
        if gb_comparison['missing_types']:
            missing_labels = [_type_label(t) for t in gb_comparison['missing_types']]
            parts.append(f'缺少国标要求的层级：{"、".join(missing_labels)}')
        gb_comparison['notes'].append('；'.join(parts))
    else:
        gb_comparison['notes'].append('编号层级与国标一致')

    # ── 逐段建议 ──────────────────────────────────────────────────────────
    suggestions = []
    for idx, ptype, match, preview in all_items:
        if ptype == 'bullet':
            suggestions.append({
                'para_index': idx,
                'numbering': match,
                'preview': preview,
                'issue': '项目符号（●◆★等）不在 GB/T 9704-2012 规定范围内，正式公文不使用',
                'suggestion': f'建议：① 改用国标第四级编号"（1）"等格式；② 或去符号直接作为正文段落',
                'severity': 'warning',
            })
        elif ptype == 'h3_multi':
            suggestions.append({
                'para_index': idx,
                'numbering': match,
                'preview': preview,
                'issue': '多级数字编号（如 1.1、4.1.1）不在 GB/T 9704-2012 规定范围内',
                'suggestion': '如果是技术报告或方案类材料可保留；如果是正式公文，建议改用"一、→（一）→1.→（1）"体系',
                'severity': 'info',
            })
        elif ptype == 'h5':
            suggestions.append({
                'para_index': idx,
                'numbering': match,
                'preview': preview,
                'issue': f'字母编号不在 GB/T 9704-2012 规定的四层序数体系中',
                'suggestion': '建议改用国标第四级编号"（1）"或圈号"①"替代',
                'severity': 'info',
            })

    # ── 层级顺序越级检查 ──────────────────────────────────────────────────
    issues = []
    _RANK_MAP = {'h1': 0, 'h2': 1, 'h3': 2, 'h3_multi': 2, 'h4': 3, 'bullet': 4}
    if seen_order:
        for current_item in all_items:
            idx, ptype, match, preview = current_item
            norm = 'h1' if ptype == 'chapter' else ('h4' if ptype == 'h5' else ptype)
            prev = None
            for prev_item in all_items:
                if prev_item[0] >= idx:
                    break
                prev = prev_item
            if prev:
                prev_norm = 'h1' if prev[1] == 'chapter' else ('h4' if prev[1] == 'h5' else prev[1])
                cur_rank = _RANK_MAP.get(norm, 99)
                prev_rank = _RANK_MAP.get(prev_norm, 99)
                if cur_rank > prev_rank + 2:
                    issues.append({
                        'para_index': idx,
                        'preview': preview,
                        'issue': f'编号 "{match}" 疑似越级：前一级是{_type_label(prev_norm)}，当前为{_type_label(norm)}',
                        'severity': 'warning',
                    })

    return {
        'levels': levels,
        'hierarchy_chain': ' → '.join(_type_label(t) for t in seen_order) if seen_order else '未检测到编号层级',
        'gb_standard': {
            'name': GB_STANDARD['name'],
            'reference': GB_STANDARD['reference'],
            'principle': GB_STANDARD['principle'],
            'description': GB_STANDARD['description'],
        },
        'gb_comparison': gb_comparison,
        'suggestions': suggestions,
        'issues': issues,
    }


def _type_label(t):
    labels = {'h1': '一级', 'h2': '二级', 'h3': '三级', 'h3_multi': '多级数字', 'h4': '四级', 'bullet': '项目符号'}
    return labels.get(t, t)


# ══════════════════════════════════════════════════════════════════════════════
# 分析模式：生成差分报告
# ══════════════════════════════════════════════════════════════════════════════

def _describe_run(run):
    f = run.font
    parts = []
    cn = '?'
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            cn = rFonts.get(qn('w:eastAsia'), '?')
        # 检查是否有加粗标记
        b = rPr.find(qn('w:b'))
        if b is not None and b.get(qn('w:val')) != 'false' and b.get(qn('w:val')) != '0':
            parts.append('加粗')
    size = str(f.size) if f.size else '?'
    parts.append(f'字体={cn}, 字号={size}')

    # 检查颜色
    if f.color and f.color.rgb and f.color.rgb != BLACK:
        parts.append(f'颜色非黑({str(f.color.rgb)})')

    return ', '.join(parts)


def _check_color(run):
    """检查 run 的字体颜色是否为黑色"""
    if run.font.color and run.font.color.rgb:
        return run.font.color.rgb != BLACK
    return False


def analyze(input_path):
    doc = Document(input_path)

    # 页面信息
    sections_info = []
    for sec in doc.sections:
        sections_info.append({
            'page_width': str(sec.page_width),
            'page_height': str(sec.page_height),
            'top_margin': str(sec.top_margin),
            'bottom_margin': str(sec.bottom_margin),
            'left_margin': str(sec.left_margin),
            'right_margin': str(sec.right_margin),
            'footer_distance': str(sec.footer_distance),
        })

    paragraphs = doc.paragraphs
    title_idx = _find_title_idx(paragraphs)

    # 颜色警告收集
    color_warnings = []
    # 加粗警告收集
    bold_warnings = []

    para_report = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        ptype, confidence, num_info = _classify(para)
        label, cn_font, size, align = TYPE_META[ptype]

        # 标题覆盖
        effective_type = ptype
        effective_confidence = confidence
        is_title_override = False
        if i == title_idx and ptype not in ('title', 'h1', 'h2', 'h3'):
            effective_type = 'title'
            effective_confidence = 'medium'
            is_title_override = True

        effective_label, effective_cn, effective_size, effective_align = TYPE_META[effective_type]

        # 当前格式
        current = {
            'alignment': str(para.alignment) if para.alignment else '?',
            'first_line_indent': str(para.paragraph_format.first_line_indent) if para.paragraph_format.first_line_indent else '0',
            'line_spacing': str(para.paragraph_format.line_spacing) if para.paragraph_format.line_spacing else '?',
        }
        run_descs = []
        has_color_issue = False
        has_bold_issue = False
        for r in para.runs:
            run_descs.append(_describe_run(r))
            if _check_color(r):
                has_color_issue = True
            if r.bold:
                has_bold_issue = True

        current['runs'] = run_descs[:3]
        if len(para.runs) > 3:
            current['runs'].append(f'... (共 {len(para.runs)} 个 run)')

        if has_color_issue:
            color_warnings.append({
                'para_index': i,
                'preview': text[:60],
                'type': effective_label,
                'message': '发现非黑色文字，政府公文正式版要求全部黑色',
            })
        if has_bold_issue:
            bold_warnings.append({
                'para_index': i,
                'preview': text[:60],
                'type': effective_label,
                'message': '当前文字有加粗，排版时将移除加粗（公文标题不要求加粗）',
            })

        # 变更列表
        changes = []
        target_align = 'CENTER' if effective_type == 'title' else 'LEFT'
        if current['alignment'] != target_align:
            changes.append(f'对齐: {current["alignment"]} → {target_align}')
        changes.append(f'中文字体 → {effective_cn} {_pt_label(effective_size)}')
        changes.append(f'西文字体 → {FONT_EN} {_pt_label(effective_size)}')
        changes.append(f'行距 → 28.9 磅')
        if has_bold_issue:
            changes.append('移除加粗 → 不加粗')
        if has_color_issue:
            changes.append('文字颜色 → 黑色')
        if effective_type == 'body' and current['first_line_indent'] != str(Pt(32)):
            changes.append('首行缩进 → 2 字符')

        # 编号信息
        num_display = None
        if num_info:
            num_display = {
                'numbering': num_info[1],
                'type': num_info[0],
            }

        para_report.append({
            'index': i,
            'preview': text[:60] + ('...' if len(text) > 60 else ''),
            'detected_as': effective_type,
            'label': effective_label,
            'confidence': effective_confidence,
            'numbering': num_display,
            'override_reason': '文档开头短文本，识别为大标题' if is_title_override else None,
            'current_format': current,
            'changes': changes,
        })

    # 表格统计
    table_count = len(doc.tables)
    table_cells = sum(len(row.cells) for table in doc.tables for row in table.rows)
    # 检查表格中文字颜色
    table_color_warnings = []
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if _check_color(run):
                            table_color_warnings.append({
                                'table_index': ti,
                                'text': para.text[:40],
                                'message': '非黑色文字',
                            })

    # 编号层级分析
    hierarchy = _analyze_numbering_hierarchy(paragraphs)

    # 封面检测
    cover_info = _detect_cover(paragraphs)

    report = {
        'source': os.path.basename(input_path),
        'sections': sections_info,
        'cover': cover_info,
        'page_changes': {
            'current': {
                'top': sections_info[0]['top_margin'] if sections_info else '?',
                'bottom': sections_info[0]['bottom_margin'] if sections_info else '?',
                'left': sections_info[0]['left_margin'] if sections_info else '?',
                'right': sections_info[0]['right_margin'] if sections_info else '?',
                'footer': sections_info[0]['footer_distance'] if sections_info else '?',
            },
            'target': {
                'top': '3.7 cm', 'bottom': '3.5 cm',
                'left': '2.7 cm', 'right': '2.7 cm',
                'footer': '2.54 cm',
                'paper': 'A4 (21.0 × 29.7 cm)',
            }
        },
        'paragraphs': para_report,
        'paragraph_summary': _summarize_paragraphs(para_report),
        'color_warnings': color_warnings,
        'bold_warnings': bold_warnings,
        'table_color_warnings': table_color_warnings,
        'numbering_hierarchy': hierarchy,
        'tables': {
            'count': table_count,
            'total_cells': table_cells,
            'change': '表格内文字统一为方正仿宋_GBK 三号（无首行缩进）' if table_count > 0 else None,
            'note': '表格表头/表体字号规范无国标强制约束，常用做法：表头黑体小四(12pt)、表体仿宋五号(10.5pt)，如有特殊要求请在确认时说明',
        },
        'page_numbers': {
            'will_add': True,
            'format': '— N —',
            'font': 'Times New Roman 四号',
            'position': '居中',
            'note': '如需精确"奇数右下+偶数左下"，排版后请用 docx 技能调整 XML',
        },
    }

    return report


def _pt_label(pt_val):
    mapping = {22: '二号', 16: '三号', 14: '四号', 12: '小四', 10.5: '五号'}
    return mapping.get(int(pt_val.pt), f'{int(pt_val.pt)}pt')


def _summarize_paragraphs(para_report):
    counts = Counter(p['detected_as'] for p in para_report)
    low_confidence = [p for p in para_report if p['confidence'] != 'high']

    type_labels = {'title': '大标题', 'h1': '一级标题', 'h2': '二级标题',
                   'h3': '三级标题', 'h4': '四级标题', 'bullet': '项目符号', 'body': '正文'}
    parts = []
    for t, label in type_labels.items():
        if counts.get(t):
            parts.append(f'{counts[t]} 个{label}')

    summary = f'共 {len(para_report)} 个有效段落（{"、".join(parts)}）'
    flags = []
    if low_confidence:
        flags.append(f'有 {len(low_confidence)} 个段落置信度较低，请人工确认')
    return {'summary': summary, 'low_confidence_count': len(low_confidence), 'warnings': flags}


# ══════════════════════════════════════════════════════════════════════════════
# 应用模式：执行排版
# ══════════════════════════════════════════════════════════════════════════════

def _blank_line():
    """创建一个空行段落（用于封面间距）"""
    return None  # 用 space_before/after 代替，不需要实体空段落


def _format_cover(doc, cover_info):
    """排版封面：标题上空行、副标题、署名区下移"""
    paragraphs = doc.paragraphs
    cover_spacing_pt = LINE_SPACING  # 28.9 pt per line

    # 标题上方空行（使用 space_before）
    for idx in cover_info['title_indices']:
        para = paragraphs[idx]
        para.paragraph_format.space_before = cover_spacing_pt * COVER_SPACE_ABOVE_TITLE

    # 署名区上方空行
    if cover_info['unit_indices']:
        for idx in cover_info['unit_indices']:
            para = paragraphs[idx]
            para.paragraph_format.space_before = cover_spacing_pt * COVER_SPACE_BELOW_TITLE

    # 封面结束后添加 page break
    if cover_info['end_index'] < len(paragraphs) - 1:
        next_para = paragraphs[cover_info['end_index'] + 1]
        # 如果下一个段落没有 page_break_before，添加
        next_para.paragraph_format.page_break_before = True

    # 格式封面字体
    for idx in cover_info['title_indices']:
        format_title(paragraphs[idx])

    for idx in cover_info.get('subtitle_indices', []):
        para = paragraphs[idx]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(para)
        _apply_font_to_runs(para, FONT_H2, FONT_EN, SIZE_3)  # 副标题用楷体三号

    for idx in cover_info.get('unit_indices', []):
        para = paragraphs[idx]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(para)
        _apply_font_to_runs(para, FONT_BODY, FONT_EN, SIZE_3)

    for idx in cover_info.get('date_indices', []):
        para = paragraphs[idx]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(para)
        _apply_font_to_runs(para, FONT_BODY, FONT_EN, SIZE_3)


def _apply_font_to_runs(para, cn_font, en_font, size):
    """对段落所有 run 设置字体，清除加粗和非黑颜色"""
    for run in para.runs:
        _set_font(run, cn_font, en_font, size)


def format_title(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(para)
    _apply_font_to_runs(para, FONT_TITLE, FONT_EN, SIZE_TITLE)


def format_h1(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para)
    _apply_font_to_runs(para, FONT_CHAPTER, FONT_EN, SIZE_CHAPTER)


def format_h2(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para)
    _apply_font_to_runs(para, FONT_SECTION, FONT_EN, SIZE_SECTION)


def format_h3(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para)
    _apply_font_to_runs(para, FONT_SUBSECTION, FONT_EN, SIZE_SUBSECTION)


def format_h4(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para)
    _apply_font_to_runs(para, FONT_ITEM, FONT_EN, SIZE_ITEM)


def format_bullet(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para)
    _apply_font_to_runs(para, FONT_BODY, FONT_EN, SIZE_BODY)


def format_body(para, indent=True):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para)
    if indent:
        para.paragraph_format.first_line_indent = Pt(int(SIZE_BODY.pt * BODY_INDENT_CHARS))
    else:
        para.paragraph_format.first_line_indent = Pt(0)
    _apply_font_to_runs(para, FONT_BODY, FONT_EN, SIZE_BODY)


_FORMATTERS = {
    'title':      format_title,
    'chapter':    format_h1,
    'section':    format_h2,
    'subsection': format_h3,
    'item':       format_h4,
    'bullet':     format_bullet,
    'body':       format_body,
    # 旧名兼容
    'h1': format_h1,
    'h2': format_h2,
    'h3': format_h3,
    'h4': format_h4,
}


def apply_format(input_path, output_path, overrides=None, structure=None):
    doc = Document(input_path)

    # 页面
    for sec in doc.sections:
        sec.top_margin = MARGINS['top']
        sec.bottom_margin = MARGINS['bottom']
        sec.left_margin = MARGINS['left']
        sec.right_margin = MARGINS['right']
        sec.footer_distance = FOOTER_DIST
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        _add_page_numbers(sec)

    # 结构映射（来自大模型判定的 structure.json）
    # structure: {"paragraphs": {idx: type}, "cover": bool|None, "title_index": int|None}
    type_map = {}
    if structure and isinstance(structure.get('paragraphs'), dict):
        type_map = {int(k): v for k, v in structure['paragraphs'].items()}
    cover_flag = structure.get('cover') if structure else None     # True/False/None
    title_index = structure.get('title_index') if structure else None

    # 兼容旧版 overrides（手动覆盖规则，优先级高于 structure）
    para_overrides = {}
    if overrides and 'paragraphs' in overrides:
        para_overrides = {int(k): v for k, v in overrides['paragraphs'].items()}

    paragraphs = doc.paragraphs

    # ── 封面检测与排版 ──
    cover_info = None
    if cover_flag is not False:
        cover_info = _detect_cover(paragraphs)
        if cover_info and cover_flag is not False:
            _format_cover(doc, cover_info)

    # 没有 structure 时，回退到启发式标题检测（兼容老流程）
    if title_index is None and not type_map:
        title_index = _find_title_idx(paragraphs)

    for i, para in enumerate(paragraphs):
        if not para.text.strip():
            continue

        # 如果封面段落已被格式化且用户未覆盖，跳过正文排版
        if cover_info and i in (cover_info['title_indices']
                                + cover_info.get('subtitle_indices', [])
                                + cover_info.get('unit_indices', [])
                                + cover_info.get('date_indices', [])):
            continue

        # 类型判定优先级：overrides > structure 映射 > 启发式
        if i in para_overrides:
            ptype = para_overrides[i]
        elif i in type_map:
            ptype = type_map[i]
        else:
            ptype, _, _ = _classify(para)
            if title_index is not None and i == title_index and ptype not in ('title', 'h1', 'h2', 'h3'):
                ptype = 'title'

        # 套用前强制清除自动编号，避免数字与正文字体不一致
        _strip_numpr(para)

        formatter = _FORMATTERS.get(ptype, format_body)
        formatter(para)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        _strip_numpr(para)
                        format_body(para, indent=False)

    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='文档排版助手 docx-helper')
    parser.add_argument('input', nargs='?', help='输入 .docx 文件')
    parser.add_argument('--reset', action='store_true', help='重置模式：清空所有手动格式，输出 {原名}+reset.docx（中性文档，便于判结构）')
    parser.add_argument('--list', action='store_true', help='列出文档所有非空段落的「索引: 文本」，用于判结构')
    parser.add_argument('--analyze', action='store_true', help='分析模式：输出 JSON 差分报告，不修改文件（启发式，仅供参考）')
    parser.add_argument('--apply', action='store_true', help='应用模式：执行排版并输出版本化文件')
    parser.add_argument('--structure', help='结构映射 JSON 文件路径（大模型判定的段落类型）')
    parser.add_argument('--config', help='排版配置文件路径（.docx-helper.json，默认查找当前目录或 ~/.workbuddy）')
    parser.add_argument('--version', type=int, help='指定版本号（默认自动递增）')
    parser.add_argument('--output', help='直接指定输出路径（跳过版本号逻辑）')
    parser.add_argument('--overrides', help='覆盖规则的 JSON 文件路径（旧版，优先级高于 structure）')

    args = parser.parse_args()

    if not args.input:
        parser.print_help()
        sys.exit(0)

    src = args.input
    if not os.path.exists(src):
        print(f'错误: 找不到 {src}')
        sys.exit(1)

    # 加载配置文件（优先级：--config > ./.docx-helper.json > ~/.workbuddy/docx-helper.json）
    config_path = args.config
    if not config_path:
        for candidate in [os.path.join(os.getcwd(), '.docx-helper.json'),
                          os.path.join(os.path.expanduser('~'), '.workbuddy', 'docx-helper.json')]:
            if os.path.exists(candidate):
                config_path = candidate
                break
    if config_path:
        print(f'配置: {config_path}')
        load_config(config_path)
    else:
        _apply_config(DEFAULT_CONFIG)  # 确保 MARGINS/TYPE_META 等被正确初始化

    if args.list:
        doc = Document(src)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f'{i}: {text}')
        sys.exit(0)

    if args.reset:
        base = os.path.splitext(os.path.basename(src))[0]
        dirname = os.path.dirname(src) or '.'
        dst = args.output or os.path.join(dirname, f'{base}+reset.docx')
        result = reset_format(src, dst)
        print(f'重置完成: {result}')
        sys.exit(0)

    if args.analyze:
        report = analyze(src)
        output = json.dumps(report, ensure_ascii=False, indent=2)
        if hasattr(sys.stdout, 'reconfigure'):
            sys.stdout.reconfigure(encoding='utf-8')
        print(output)
        sys.exit(0)

    if args.output:
        dst = args.output
    else:
        base, version, dst = _next_version(src)
        if args.version:
            version = args.version
            dirname = os.path.dirname(src) or '.'
            dst = os.path.join(dirname, f'{base}+docx-helper+v{version}.docx')
        print(f'版本: v{version}')

    structure = None
    if args.structure and os.path.exists(args.structure):
        with open(args.structure, 'r', encoding='utf-8') as f:
            structure = json.load(f)

    overrides = None
    if args.overrides and os.path.exists(args.overrides):
        with open(args.overrides, 'r', encoding='utf-8') as f:
            overrides = json.load(f)

    result = apply_format(src, dst, overrides, structure)
    print(f'完成: {result}')
